import imaplib
import email
from email.header import decode_header
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader
import pytesseract
from pdf2image import convert_from_path
import spacy
from word2number import w2n
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import re
import io
import os
from dotenv import load_dotenv
from rapidfuzz import fuzz, process  
import comtypes.client

# Load spaCy's transformer-based model for better NER accuracy
nlp = spacy.load("en_core_web_trf")

# Path to Tesseract-OCR executable (adjust for your system)
pytesseract.pytesseract.tesseract_cmd = r"C:\Users\andre\anaconda3\Lib\site-packages\tesseract"

# Load environment variables
env_path = r"C:\Users\Lucas\OneDrive\Desktop\New folder\credentials.env"
load_dotenv(env_path)

# Sanitize filenames (remove problematic characters)
def sanitize_filename(filename):
    sanitized = re.sub(r'[<>:"/\\|?*]', '_', filename)
    sanitized = sanitized.replace('\r', '').replace('\n', '')  # Remove newlines
    sanitized = sanitized.replace('\t', '')  # Remove tabs
    return sanitized

# Extract email body
def extract_email_body(msg):
    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            content_disposition = str(part.get("Content-Disposition"))
            if "attachment" not in content_disposition:
                if content_type == "text/plain":
                    return part.get_payload(decode=True).decode("utf-8", errors="ignore")
                elif content_type == "text/html":
                    return part.get_payload(decode=True).decode("utf-8", errors="ignore")
    else:
        return msg.get_payload(decode=True).decode("utf-8", errors="ignore")
    

# Extract text from DOCX
def extract_text_from_docx(attachment_content):
    doc = Document(io.BytesIO(attachment_content))
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Extract text from PDF
def extract_text_from_pdf(attachment_content):
    pdf_reader = PdfReader(io.BytesIO(attachment_content))
    text = ""
    for page in range(len(pdf_reader.pages)):
        text += pdf_reader.pages[page].extract_text()
    return text

# OCR for scanned PDFs
def extract_text_with_ocr(pdf_path):
    text = ""
    try:
        images = convert_from_path(pdf_path)
        for image in images:
            text += pytesseract.image_to_string(image, lang='eng')
    except Exception as e:
        print(f"OCR failed for {pdf_path}: {e}")
    return text

# Read resume files
def read_resume_from_file(file_path):
    text = ""
    try:
        if file_path.endswith(".docx"):
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif file_path.endswith(".pdf"):
            reader = PdfReader(file_path)
            text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
            if not text.strip():
                text = extract_text_with_ocr(file_path)
        elif file_path.endswith(".doc"):

            word = comtypes.client.CreateObject('Word.Application')
            doc = word.Documents.Open(file_path)
            text = doc.Content.Text
            doc.Close()
            word.Quit()
        else:
            raise ValueError("Unsupported file format")
    except Exception as e:
        print(f"Error reading file {file_path}: {e}")
    return text

# Extract resume details
def extract_name_from_text(text):
    text = text.strip()
    text = re.sub(r'\S+@\S+', '', text)
    text = re.sub(r'[^a-zA-Z\s]', '', text)
    lines = text.split("\n")
    irrelevant_words = ["summary", "contact", "education", "experience", "skills", "references", "profile", "resume", "cv"]
    for line in lines[:3]:
        line = line.strip()
        if any(irrelevant_word in line.lower() for irrelevant_word in irrelevant_words):
            continue
        if len(line) > 1:
            name_parts = line.split()
            if len(name_parts) > 1:
                return " ".join([part.title() for part in name_parts])
            elif len(name_parts) == 1:
                return name_parts[0].title()
    return "Name not found"

# Function to extract email from resume text
def extract_email_from_text(text):
    email_match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return email_match.group(0) if email_match else "Email not found"

# Function to extract phone numbers
def extract_phone_from_text(text):
    phone_pattern = re.compile(r"(?:direct|mobile|phone|ph#|contact|tel|cell)?[:\s-]*"
                               r"(?:\+?\d{1,3}[-.\s]?)?"
                               r"\(?\d{1,4}\)?"
                               r"[-.\s]?\d{1,4}"
                               r"[-.\s]?\d{1,4}"
                               r"[-.\s]?\d{1,9}"
                               r"(?:\s?(?:ext|x|extension)\s?\d{1,5})?")
    matches = phone_pattern.findall(text)
    phones = [re.sub(r"[^+\d\s()-]", "", match).strip() for match in matches if len(re.sub(r"\D", "", match)) >= 10]
    return ", ".join(phones) if phones else "Phone not found"

# Function to extract experience from resume text
def extract_experience(text):
    text = text.lower()
    numeric_pattern = r"(?:more than|over|at least|around|approximately|nearly|up to)?\s*(\d+)\+?\s*years?"
    numeric_match = re.search(numeric_pattern, text)
    if numeric_match:
        years = numeric_match.group(1)
        return f"{int(years)}+ years" if '+' in numeric_match.group(0) else f"{int(years)} years"
    return "Experience not found"

# Function to extract skills from job description
def extract_skills_from_job_description(job_desc_text):
    """
    Extract skills from the job description text.
    Assumes skills are mentioned after keywords like 'with' or 'experience'.
    """
    # Extract skills after 'with' or 'experience'
    skill_pattern = re.compile(r'(?:with|experience)\s+([\w\s,/#+]+)', re.IGNORECASE)
    matches = skill_pattern.findall(job_desc_text)
    
    # Split skills by commas, slashes, or the word 'and'
    skills = []
    for match in matches:
        # Split by commas, slashes, or 'and'
        skills.extend(re.split(r'[,/]| and ', match))
    
    # Clean and filter skills
    cleaned_skills = []
    for skill in skills:
        skill = skill.strip()
        if skill and len(skill) > 1:  # Filter out empty or single-character strings
            # Remove any trailing or leading non-alphanumeric characters
            skill = re.sub(r'^\W+|\W+$', '', skill)
            # Exclude irrelevant phrases like "experience", "Location", etc.
            if not any(irrelevant_word in skill.lower() for irrelevant_word in ["experience", "location", "scope", "data", "security", "document", "purchase", "order", "application", "requirements", "best", "practice", "design", "developing", "functions", "including", "tables", "interpreting", "specifications", "reading", "stored", "procedures", "test", "cases", "triggers", "views", "writing", "zero"]):
                cleaned_skills.append(skill)
    
    # Remove duplicates and sort
    unique_skills = sorted(set(cleaned_skills), key=lambda x: x.lower())
    
    return unique_skills

# Predefined skill set (expand as needed)
def extract_relevant_skills(resume_text, job_desc_subject):
    """
    Extract skills from job description subject and check if they exist in the resume.
    The subject will only mention skills after specific keywords like 'with'.
    """
    # Extract skills after 'with'
    skill_pattern = re.compile(r'with (.+)', re.IGNORECASE)
    match = skill_pattern.search(job_desc_subject)
    
    if not match:
        return []
    
    # Get skills list from subject
    skills_str = match.group(1).strip()
    subject_skills = [skill.strip().lower() for skill in re.split(r',|/', skills_str)]
    
    # Prepare resume text for matching
    resume_lower = resume_text.lower()
    
    # Find matches with whole word matching
    matched_skills = []
    for skill in subject_skills:
        pattern = re.compile(rf'\b{re.escape(skill)}\b', re.IGNORECASE)
        if pattern.search(resume_lower):
            matched_skills.append(skill.title())
    
    return matched_skills

# Function to extract certifications
def extract_certifications_count(text):
    certification_keywords = [
        r"certification", r"certifications", r"certified", r"certificate", r"certificates"
    ]
    pattern = r"|".join(certification_keywords)
    matches = re.findall(pattern, text, re.IGNORECASE)
    return len(matches)

# Extract government column from resume text
def extract_government_details(resume_text):
    """
    Extracts U.S. government-related information from resume text.
    Uses a fully merged and optimized regex pattern set to identify U.S. government agencies and departments.
    """

    # Comprehensive regex patterns for U.S. government entities
    patterns = [
        # Full department and agency names
    r"\b(?:U\.S\.|United States|Federal|National|State of|City of) [A-Za-z\s]+ (?:Department|Agency|Bureau|Commission|Administration|Office|Authority|Service)\b",
    r"\b(?:Department|Office|Agency|Bureau|Commission|Administration|Authority) of [A-Za-z\s]+\b",

    # Broad U.S. government-related phrases
    r"\b(?:U\.S\.|United States|Federal|National) [A-Za-z\s]+ (?:Service|Program|Initiative)\b",
    r"\b(?:Agency for|Bureau of|Commission on|Administration for|Office of) [A-Za-z\s]+\b",

    # Government branches and major entities
    r"\b(?:Congress|Senate|House of Representatives|White House|Supreme Court|Judicial Branch|Legislative Branch|Executive Branch)\b",

    # Military branches
    r"\b(?:Air Force|Army|Navy|Marine Corps|Coast Guard)\b",

    # Specific agencies and organizations
    r"\b(?:National Security Agency|Central Intelligence Agency|Federal Bureau of Investigation|Internal Revenue Service|Environmental Protection Agency|Department of Defense|Federal Reserve|Department of Justice|Department of Commerce|Treasury Department|Veterans Affairs|Social Security Administration|Centers for Medicare & Medicaid Services|General Services Administration|Federal Aviation Administration|Federal Communications Commission)\b",

    # Common U.S. government agency acronyms
    r"\b(?:USPS|USCIS|ICE|TSA|FBI|CIA|NSA|DIA|GSA|CMS|EPA|FAA|FCC|VA|NASA|NIH|NOAA|NPS|NSF|NTSB|NRC|OPM|SBA|USAID|USGS|USMS|USSS|DARPA|NGA|NRO|ODNI|FEMA|CDC|FDA|HRSA|IHS|SAMHSA|OSHA|EEOC|DOE|NREL|DOT|FHA|FRA|FTA|MARAD|HUD|USDA|FNS|FS|NRCS|IRS|OCC|FinCEN|EXIM|NIST)\b",

    # Homeland Security and law enforcement
    r"\b(?:Homeland Security|Customs and Border Protection|Defense Intelligence Agency)\b",

    # Security Clearance and classified projects
    r"\b(?:Security Clearance|Top Secret|Classified Project)\b",

    # Specific programs and initiatives
    r"\b(?:Affordable Care Act|Medicare|Medicaid|Social Security|Veterans Health Administration|Federal Student Aid|National Flood Insurance Program|Supplemental Nutrition Assistance Program|Temporary Assistance for Needy Families)\b",

    # Emergency and disaster management
    r"\b(?:Federal Emergency Management Agency|National Disaster Recovery Framework|National Response Framework)\b",

    # Health and human services
    r"\b(?:Centers for Disease Control and Prevention|Food and Drug Administration|Health Resources and Services Administration|Indian Health Service|Substance Abuse and Mental Health Services Administration)\b",

    # Education and labor
    r"\b(?:Department of Education|Department of Labor|Occupational Safety and Health Administration|Equal Employment Opportunity Commission)\b",

    # Energy and environment
    r"\b(?:Department of Energy|Nuclear Energy Commission|National Renewable Energy Laboratory|Environmental Protection Agency)\b",

    # Transportation and infrastructure
    r"\b(?:Department of Transportation|Federal Highway Administration|Federal Railroad Administration|Federal Transit Administration|Maritime Administration)\b",

    # Housing and urban development
    r"\b(?:Department of Housing and Urban Development|Federal Housing Administration|Office of Public and Indian Housing)\b",

    # Agriculture and food
    r"\b(?:Department of Agriculture|Food and Nutrition Service|Forest Service|Natural Resources Conservation Service)\b",

    # Treasury and finance
    r"\b(?:Department of the Treasury|Internal Revenue Service|Office of the Comptroller of the Currency|Financial Crimes Enforcement Network)\b",

    # International affairs
    r"\b(?:Department of State|U.S. Agency for International Development|Peace Corps|Export-Import Bank of the United States)\b",

    # Veterans and military affairs
    r"\b(?:Department of Veterans Affairs|Veterans Health Administration|National Cemetery Administration)\b",

    # Science and technology
    r"\b(?:National Aeronautics and Space Administration|National Institute of Standards and Technology|National Science Foundation)\b",
    ]

    # Collect matched government details
    government_details = set()

    for pattern in patterns:
        matches = re.findall(pattern, resume_text, re.IGNORECASE)
        for match in matches:
            clean_match = match.strip()
            if clean_match.lower() not in [
                "public sector", "government contractor", "government client",
                "administration of weblogic environment", "state of art technologies"
            ]:
                government_details.add(clean_match)

    # If no valid government organizations are found, return only "Not Worked with Government"
    return "Not Worked with Government" if not government_details else ", ".join(sorted(government_details))

# Function to extract location from resume text
def extract_location_from_text(text):
    """Extract location (city, state, or ZIP code) from resume text."""
    location_match = re.search(
        r"\b([A-Z][a-z]+(?:\s[A-Z][a-z]+)*,\s(?:TX|CA|NY|FL|WA|IL|PA|GA|NC|OH|NJ|VA|CO|AZ|MA|MD|TN|MO|IN|WI|MN|SC|AL|LA|KY|OR|OK|CT|IA|MS|KS|AR|NV|UT|NM|NE|WV|ID|HI|ME|NH|MT|RI|DE|SD|ND|AK|VT|WY))\b"  # City, State
        r"|\b\d{5}(?:-\d{4})?\b",  # ZIP code
        text
    )
    if location_match:
        location = location_match.group(0)
        if not any(keyword in location.lower() for keyword in ["assistant", "server", "sql"]):  # Example of filtering out unrelated matches
            return location
    return "Location not found"

# Function to extract visa status from the resume text
def extract_visa_status(text):
    """Extract visa status from the resume text."""
    visa_keywords = {
        "H1B": ["h1b"],
        "Green Card": ["green card", "permanent resident"],
        "US Citizen": ["usc", "us citizen", "citizenship: us"],
        "OPT": ["opt"],
        "CPT": ["cpt"],
        "L2": ["l2 visa"],
        "EAD": ["ead"],
        "TN Visa": ["tn visa"],
        "Study Visa": ["study visa"]
    }
    visa_status = []
    for visa, patterns in visa_keywords.items():
        for pattern in patterns:
            if re.search(pattern, text.lower()):
                visa_status.append(visa)
                break
    return ", ".join(visa_status) if visa_status else "Not found"

# Calculating resume score
def calculate_resume_score(resume_text, job_desc_text, skills, experience, certifications, visa_status, location, government):
    corpus = [job_desc_text, resume_text]
    vectorizer = CountVectorizer().fit_transform(corpus)
    vectors = vectorizer.toarray()

    # Cosine Similarity: Measures how closely the resume text aligns with the job description.
    similarity_score = cosine_similarity([vectors[0]], [vectors[1]])[0][0]

    # Skills, experience, and certifications
    skills_count = len(skills)
    experience_years = int(re.search(r"\d+", experience).group(0)) if re.search(r"\d+", experience) else 0
    certifications_count = certifications
 
    normalized_experience = min(experience_years / 20, 1)
    normalized_skills = min(skills_count / 20, 1)
 
    # Visa Status Scoring
    visa_priority = {
        "US Citizen": 1.0,
        "Green Card": 0.9,
        "H1B": 0.8,
        "OPT": 0.7,
        "CPT": 0.6,
        "L2": 0.5,
        "EAD": 0.5,
        "TN Visa": 0.6,
        "Not found": 0.0
    }
    visa_score = visa_priority.get(visa_status, 0.0)
 
    # Location Scoring
    location_score = 0.0
    if location.lower() != "location not found":
        location_score = 1.0  
 
    # Government Scoring
    government_score = 0.0
    government_str = ", ".join(government) if government else "Not found"
    if government_str.lower() != "not found":
        government_score = 1.0 


    # Weighted scoring
    score = (
        similarity_score * 0.5 +           # Adjusted to 50% weight
        normalized_skills * 0.8 +          # Adjusted to 80% weight
        normalized_experience * 0.01 +     # Adjusted to 1% weight
        certifications_count * 0.01 +      # Certifications contribute 1%
        visa_score * 0.05 +                # Visa status contributes 5%
        location_score * 0.05 +            # Location contributes 5%
        government_score * 0.05            # Government experience contributes 5%
    )

    return round(min(score * 100, 100), 2)

# Function to filter emails by Job ID in subject or body
def filter_emails_by_job_id(job_id, email_ids, mail):
    filtered_emails = []
    for email_id in email_ids:
        status, msg_data = mail.fetch(email_id, "(RFC822)")
        for response_part in msg_data:
            if isinstance(response_part, tuple):
                msg = email.message_from_bytes(response_part[1])
                subject, encoding = decode_header(msg["Subject"])[0]
                if isinstance(subject, bytes):
                    subject = subject.decode(encoding or "utf-8")
                # Match Job ID in Subject and Body (case insensitive)
                if job_id.lower() in subject.lower():
                    filtered_emails.append(msg)
                else:
                    body = extract_email_body(msg)
                    if body and job_id.lower() in body.lower():
                        filtered_emails.append(msg)
    return filtered_emails

# Process resumes and filter emails based on Job ID
def process_resumes_and_attachments(job_id):

    try:
        # Get email credentials from environment variables
        email_user = os.getenv('EMAIL_USERNAME')
        email_pass = os.getenv('EMAIL_PASSWORD')

        mail = imaplib.IMAP4_SSL('imap.gmail.com')
        mail.login(email_user, email_pass)
        mail.select("inbox")

        print(f"Processing emails for Job ID: {job_id}")
        status, messages = mail.search(None, 'ALL')
        email_ids = messages[0].split()

        # Filter emails by Job ID
        filtered_emails = filter_emails_by_job_id(job_id, email_ids, mail)

        # Check if no emails match the Job ID
        if not filtered_emails:
            raise ValueError(f"No emails found for Job ID: {job_id}. Please provide a valid Job ID.")
        
        print(f"Found {len(filtered_emails)} emails matching the Job ID: {job_id}")

        resume_details = []
        job_desc_skills = set()  
        
        # Create a folder to store resumes with "Name not found" for the current job ID
        output_folder = f"Unread_Resumes"
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
            """print(f"Created folder: {output_folder}")"""
        else:
            # Clear the folder if it already exists (optional)
            for file in os.listdir(output_folder):
                file_path = os.path.join(output_folder, file)
                try:
                    if os.path.isfile(file_path):
                        os.unlink(file_path)
                except Exception as e:
                    print(f"Failed to delete {file_path}: {e}")
            '''print(f"Cleared existing folder: {output_folder}")'''

        # List to track resumes with "Name not found"
        resumes_without_names = []

        for msg in filtered_emails:
            job_desc_text = extract_email_body(msg)

            # Extract skills from the job description
            job_desc_skills.update(extract_skills_from_job_description(job_desc_text))

            for part in msg.walk():
                if part.get("Content-Disposition") and "attachment" in part.get("Content-Disposition"):
                    attachment_filename = sanitize_filename(part.get_filename())
                    attachment_content = part.get_payload(decode=True)

                    if attachment_filename.lower().endswith('.pdf'):
                        resume_text = extract_text_from_pdf(attachment_content)
                    elif attachment_filename.lower().endswith('.docx'):
                        resume_text = extract_text_from_docx(attachment_content)
                    else:
                        continue

                    # Extract details from the resume text
                    details = {
                        "name": extract_name_from_text(resume_text),
                        "email": extract_email_from_text(resume_text),
                        "phone": extract_phone_from_text(resume_text),
                        "experience": extract_experience(resume_text),
                        "skills": extract_relevant_skills(resume_text, job_desc_text),
                        "certifications": extract_certifications_count(resume_text),
                        "location": extract_location_from_text(resume_text),
                        "visa_status": extract_visa_status(resume_text),
                        "government": extract_government_details(resume_text)
                    }

                    # Calculate the resume score
                    score = calculate_resume_score(resume_text, job_desc_text, details['skills'],
                                                details['experience'], details['certifications'],
                                                details['visa_status'], details['location'], details['government'])
                    details['Resume Score'] = score

                    # Save resumes with "Name not found" to the separate folder
                    if details['name'] == "Name not found":
                        resume_path = os.path.join(output_folder, attachment_filename)
                        with open(resume_path, "wb") as resume_file:
                            resume_file.write(attachment_content)
                        '''print(f"Saved resume with missing name: {resume_path}")'''
                        resumes_without_names.append(attachment_filename)  # Track these resumes
                    else:
                        resume_details.append(details)  # Only add resumes with valid names to the table

        mail.logout()

        # Create DataFrame from the extracted resume details
        df = pd.DataFrame(resume_details)

        # Create rank function using resume score
        def assign_rank(score):
            score = int(score)  # Convert the score to an integer
            if 0 <= score <= 9:
                return 10
            elif 10 <= score <= 19:
                return 9
            elif 20 <= score <= 29:
                return 8
            elif 30 <= score <= 39:
                return 7
            elif 40 <= score <= 49:
                return 6
            elif 50 <= score <= 59:
                return 5
            elif 60 <= score <= 69:
                return 4
            elif 70 <= score <= 79:
                return 3
            elif 80 <= score <= 89:
                return 2
            elif 90 <= score <= 100:
                return 1
            return 10  # Default if something goes wrong

        # Apply rank assignment based on the score
        df['Rank'] = df['Resume Score'].apply(assign_rank)  # Assign ranks
        df = df.sort_values(by="Rank", ascending=True).reset_index(drop=True)  # Sort in ascending order

        '''
        # Print the list of resumes with missing names
       
        #if resumes_without_names:
        #print(f"Resumes with missing names saved in '{output_folder}':")
        #for resume in resumes_without_names:
        #print(f"- {resume}")
                                    '''
        
        # Display the skills mentioned in the job description
        print("\nJob Description Skills:")
        print(", ".join(sorted(job_desc_skills)))  # Sort and display skills as a comma-separated list

        return df


    except ValueError as ve:
        print(f"Error: {ve}")
        return None  # Return None to indicate no emails were found
    except imaplib.IMAP4.error as e:
        print(f"Failed to connect to the email server: {e}")
        return None
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return None

# Example usage:
job_id = input("Enter the Job ID to search: ")
df = process_resumes_and_attachments(job_id)

